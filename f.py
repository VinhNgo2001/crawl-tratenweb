import requests
from bs4 import BeautifulSoup
import re
import json
import base64
from io import BytesIO
from docx import Document
from docx.shared import Inches
import os

# Các field cần lấy (trừ Mã số thuế, SĐT sẽ xử lý riêng)
DETAIL_FIELDS = {
    "Tên giao dịch": r"Tên giao dịch:\s*(.+)",
    "Địa chỉ": r"Địa chỉ:\s*(.+)",
    "Đại diện pháp luật": r"Đại diện pháp luật:\s*(.+)",
    "Ngày cấp giấy phép": r"Ngày cấp giấy phép:\s*(.+)",
    "Ngày hoạt động": r"Ngày hoạt động:\s*(.+)",
    "Trạng thái": r"Trạng thái:\s*(.+)"
}

def add_base64_image(doc, base64_str, width=1.2):
    """Giải mã ảnh base64 và chèn vào docx"""
    try:
        if base64_str.startswith("data:image"):
            base64_str = base64_str.split(",")[1]
        img_data = base64.b64decode(base64_str)
        image_stream = BytesIO(img_data)
        doc.add_picture(image_stream, width=Inches(width))
    except Exception as e:
        doc.add_paragraph(f"[Lỗi ảnh: {e}]")

def get_company_detail(url):
    res = requests.get(url)
    soup = BeautifulSoup(res.text, "html.parser")
    
    info = {"url": url, "soup": soup}
    block = soup.select_one("div.jumbotron")
    if block:
        text = block.get_text("\n", strip=True)

        # lấy các field text
        for field, pattern in DETAIL_FIELDS.items():
            match = re.search(pattern, text)
            if match:
                info[field] = match.group(1).strip()

    return info

def next_docx_name(base_name, folder, ext=".docx"):
    """
    Sinh tên file theo bộ đếm tăng dần: base_name_1.docx, base_name_2.docx, ...
    """
    files = os.listdir(folder)
    numbers = []
    for f in files:
        if f.startswith(base_name) and f.endswith(ext):
            parts = f.replace(ext, "").split("_")
            if len(parts) > 1 and parts[-1].isdigit():
                numbers.append(int(parts[-1]))
    next_num = max(numbers) + 1 if numbers else 1
    return os.path.join(folder, f"{base_name}_{next_num}{ext}")

if __name__ == "__main__":
    # Đọc danh sách link công ty
    with open("companies_links_page_12_13.json", "r", encoding="utf-8") as f:
        companies = json.load(f)
    
    # Tạo file Word
    doc = Document()

    for comp in companies:
        if comp["ten_cong_ty"].strip().upper().startswith("UBND"):
            continue
        print("📌 Đang lấy:", comp["ten_cong_ty"])
        detail = get_company_detail(comp["url"])
        soup = detail["soup"]

        # --- Ghi vào Word ---
        doc.add_heading(comp["ten_cong_ty"], level=1)

        if "Địa chỉ" in detail:
            doc.add_paragraph(f"Địa chỉ: {detail['Địa chỉ']}")
        if "Ngày hoạt động" in detail:
            doc.add_paragraph(f"Ngày hoạt động: {detail['Ngày hoạt động']}")
        if "Đại diện pháp luật" in detail:
            doc.add_paragraph(f"Đại diện pháp luật: {detail['Đại diện pháp luật']}")

        # xử lý ảnh base64 (thường có 2 ảnh: mã số thuế, số điện thoại)
        imgs = soup.select("img")

        # --- Mã số thuế ---
        if imgs and len(imgs) > 0 and imgs[0].get("src", "").startswith("data:image"):
            doc.add_paragraph("Mã số thuế: ")
            add_base64_image(doc, imgs[0]["src"], width=1.2)
        elif "Mã số thuế" in detail:
            doc.add_paragraph(f"Mã số thuế: {detail['Mã số thuế']}")

        # --- Số điện thoại ---
        if imgs and len(imgs) > 1 and imgs[1].get("src", "").startswith("data:image"):
            doc.add_paragraph("Số điện thoại: ")
            add_base64_image(doc, imgs[1]["src"], width=1.0)
        elif "Điện thoại trụ sở" in detail:
            doc.add_paragraph(f"Số điện thoại: {detail['Điện thoại trụ sở']}")

        doc.add_paragraph("-------------------------")

    # Xuất ra file với tên tự tăng
    output_folder = "."   # thư mục hiện tại, bạn có thể đổi thành "D:/output"
    file_path = next_docx_name("companies_output", output_folder)
    doc.save(file_path)
    print(f"✅ Đã xuất ra {file_path}")
