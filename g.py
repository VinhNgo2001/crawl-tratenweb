import requests
from bs4 import BeautifulSoup
import re
import json
import base64
from io import BytesIO
from docx import Document
from docx.shared import Inches
import os
import time
import random
from concurrent.futures import ThreadPoolExecutor, as_completed

# Headers giả lập trình duyệt
HEADERS = {
    "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
                  "AppleWebKit/537.36 (KHTML, like Gecko) "
                  "Chrome/120.0 Safari/537.36",
    "Accept-Language": "vi,en;q=0.9",
}

# Các field cần lấy
DETAIL_FIELDS = {
    "Tên giao dịch": r"Tên giao dịch:\s*(.+)",
    "Địa chỉ": r"Địa chỉ:\s*(.+)",
    "Đại diện pháp luật": r"Đại diện pháp luật:\s*(.+)",
    "Ngày cấp giấy phép": r"Ngày cấp giấy phép:\s*(.+)",
    "Ngày hoạt động": r"Ngày hoạt động:\s*(.+)",
    "Trạng thái": r"Trạng thái:\s*(.+)"
}

def add_base64_image(doc, base64_str, width=1.2):
    """Giải mã ảnh base64 và chèn vào docx"""
    try:
        if base64_str.startswith("data:image"):
            base64_str = base64_str.split(",")[1]
        img_data = base64.b64decode(base64_str)
        image_stream = BytesIO(img_data)
        doc.add_picture(image_stream, width=Inches(width))
    except Exception as e:
        doc.add_paragraph(f"[Lỗi ảnh: {e}]")

def get_company_detail(url, retries=3, delay=3):
    """Crawl chi tiết công ty với retry"""
    for attempt in range(retries):
        try:
            res = requests.get(url, headers=HEADERS, timeout=15)
            if res.status_code != 200:
                raise Exception(f"HTTP {res.status_code}")
            soup = BeautifulSoup(res.text, "html.parser")

            info = {"url": url, "soup": soup}
            block = soup.select_one("div.jumbotron")
            if block:
                text = block.get_text("\n", strip=True)
                for field, pattern in DETAIL_FIELDS.items():
                    match = re.search(pattern, text)
                    if match:
                        info[field] = match.group(1).strip()
            return info

        except Exception as e:
            print(f"⚠️ Lỗi {url}: {e} (lần thử {attempt+1})")
            time.sleep(delay + random.random())  # chờ rồi thử lại

    return {"url": url, "error": True, "soup": None}

def next_docx_name(base_name, folder, ext=".docx"):
    """Sinh tên file theo bộ đếm tăng dần"""
    files = os.listdir(folder)
    numbers = []
    for f in files:
        if f.startswith(base_name) and f.endswith(ext):
            parts = f.replace(ext, "").split("_")
            if len(parts) > 1 and parts[-1].isdigit():
                numbers.append(int(parts[-1]))
    next_num = max(numbers) + 1 if numbers else 1
    return os.path.join(folder, f"{base_name}_{next_num}{ext}")

if __name__ == "__main__":
    # Đọc danh sách link công ty
    with open("companies_links_page_12_22.json", "r", encoding="utf-8") as f:
        companies = json.load(f)

    doc = Document()

    results = []
    # Crawl song song 10 luồng
    with ThreadPoolExecutor(max_workers=10) as executor:
        future_to_comp = {
            executor.submit(get_company_detail, comp["url"]): comp
            for comp in companies
            if not comp["ten_cong_ty"].strip().upper().startswith("UBND")
        }

        for future in as_completed(future_to_comp):
            comp = future_to_comp[future]
            try:
                detail = future.result()
                results.append((comp, detail))
            except Exception as e:
                print(f"❌ Lỗi khi crawl {comp['ten_cong_ty']}: {e}")

    # Ghi kết quả vào Word
    for comp, detail in results:
        soup = detail.get("soup")
        doc.add_heading(comp["ten_cong_ty"], level=1)

        if "Địa chỉ" in detail:
            doc.add_paragraph(f"Địa chỉ: {detail['Địa chỉ']}")
        if "Ngày hoạt động" in detail:
            doc.add_paragraph(f"Ngày hoạt động: {detail['Ngày hoạt động']}")
        if "Đại diện pháp luật" in detail:
            doc.add_paragraph(f"Đại diện pháp luật: {detail['Đại diện pháp luật']}")

        # ảnh base64 (MST + SĐT)
        if soup:
            imgs = soup.select("img")

            if imgs and len(imgs) > 0 and imgs[0].get("src", "").startswith("data:image"):
                doc.add_paragraph("Mã số thuế: ")
                add_base64_image(doc, imgs[0]["src"], width=1.2)

            if imgs and len(imgs) > 1 and imgs[1].get("src", "").startswith("data:image"):
                doc.add_paragraph("Số điện thoại: ")
                add_base64_image(doc, imgs[1]["src"], width=1.0)

        doc.add_paragraph("-------------------------")

    # Xuất file
    output_folder = "."
    file_path = next_docx_name("companies_output", output_folder)
    doc.save(file_path)
    print(f"✅ Đã xuất ra {file_path}")
