import requests
from bs4 import BeautifulSoup
import re
import json
import base64
from io import BytesIO
from docx import Document
from docx.shared import Inches

# Các field cần lấy
DETAIL_FIELDS = {
    "Tên giao dịch": r"Tên giao dịch:\s*(.+)",
    "Mã số thuế": r"Mã số thuế:\s*(.+)",
    "Địa chỉ": r"Địa chỉ:\s*(.+)",
    "Đại diện pháp luật": r"Đại diện pháp luật:\s*(.+)",
    "Ngày cấp giấy phép": r"Ngày cấp giấy phép:\s*(.+)",
    "Ngày hoạt động": r"Ngày hoạt động:\s*(.+)",
    "Điện thoại trụ sở": r"Điện thoại trụ sở:\s*(.+)",
    "Trạng thái": r"Trạng thái:\s*(.+)"
}

def add_base64_image(doc, base64_str, width=0.8):
    """Giải mã ảnh base64 và chèn vào docx với size nhỏ"""
    try:
        if base64_str.startswith("data:image"):
            base64_str = base64_str.split(",")[1]
        img_data = base64.b64decode(base64_str)
        image_stream = BytesIO(img_data)
        doc.add_picture(image_stream, width=Inches(width))
    except Exception as e:
        doc.add_paragraph(f"[Lỗi ảnh: {e}]")

def get_company_detail(url):
    res = requests.get(url)
    soup = BeautifulSoup(res.text, "html.parser")
    
    info = {"url": url, "soup": soup}
    block = soup.select_one("div.jumbotron")
    if block:
        text = block.get_text("\n", strip=True)

        for field, pattern in DETAIL_FIELDS.items():
            match = re.search(pattern, text)
            if match:
                info[field] = match.group(1).strip()
    return info

if __name__ == "__main__":
    # Đọc danh sách link công ty
    with open("companies_links_page_12_22.json", "r", encoding="utf-8") as f:
        companies = json.load(f)
    
    # Tạo file Word
    doc = Document()

    for comp in companies:
        if comp["ten_cong_ty"].strip().upper().startswith("UBND"):
         continue
        print("📌 Đang lấy:", comp["ten_cong_ty"])
        detail = get_company_detail(comp["url"])

        # --- Ghi vào Word ---
        doc.add_heading(comp["ten_cong_ty"], level=1)

        if "Địa chỉ" in detail:
            doc.add_paragraph(f"Địa chỉ: {detail['Địa chỉ']}")
        if "Ngày hoạt động" in detail:
            doc.add_paragraph(f"Ngày hoạt động: {detail['Ngày hoạt động']}")
        if "Đại diện pháp luật" in detail:
            doc.add_paragraph(f"Đại diện pháp luật: {detail['Đại diện pháp luật']}")

        # Xử lý điện thoại: nếu có ảnh base64 thì chèn ảnh, nếu text thì in text
        phone_img = detail["soup"].select_one("img")
        if phone_img and phone_img.get("src", "").startswith("data:image"):
            p = doc.add_paragraph("Số điện thoại: ")
            add_base64_image(doc, phone_img["src"], width=0.8)
        elif "Điện thoại trụ sở" in detail:
            doc.add_paragraph(f"Số điện thoại: {detail['Điện thoại trụ sở']}")

        doc.add_paragraph("-------------------------")

    # Xuất ra file
    doc.save("companies_details_11_22.docx")
    print("✅ Đã xuất ra companies_details2.docx")
